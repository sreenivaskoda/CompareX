"""
Enhanced DOCX Document Comparison Tool
A Flask-based application for comparing Microsoft Word documents with advanced features:
- Text content comparison with difflib
- Strikethrough detection and tracking
- Table structure and content analysis
- Image extraction and comparison
- Excel export with comprehensive reporting
- Encrypted document support
- NEW: Machine learning-based change classification
- NEW: Change severity scoring
- NEW: Document similarity analysis
- NEW: Advanced filtering and search
- NEW: Batch processing support
- NEW: Real-time progress tracking
- NEW: Advanced visualization options
"""

import os
import io
import tempfile
import difflib
import re
import json
import time
import threading
from hashlib import md5
from collections import defaultdict, Counter
from typing import List, Dict, Any, Tuple, Optional, Set
from datetime import datetime
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed

# Third-party imports
from docx.shared import RGBColor
import msoffcrypto
import docx
import openpyxl
from flask import Flask, render_template, request, send_file, flash, jsonify, redirect, url_for, session
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from openpyxl.chart import PieChart, Reference, BarChart
from PIL import Image as PILImage
import xml.etree.ElementTree as ET
from functools import lru_cache
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import DBSCAN
import seaborn as sns
import matplotlib.pyplot as plt
import base64

# ================================
# APPLICATION CONFIGURATION
# ================================

app = Flask(__name__)
app.secret_key = "supersecretkey"  # Change this in production
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# File upload configuration
UPLOAD_FOLDER = "uploads"
COMPARISON_CACHE = "cache"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(COMPARISON_CACHE, exist_ok=True)

# Global variables for progress tracking
progress_data = {}

# ================================
# NEW: Machine Learning Components
# ================================

class ChangeClassifier:
    """Machine learning-based change classification system"""
    
    def __init__(self):
        self.change_categories = {
            'formatting': ['font', 'size', 'color', 'bold', 'italic', 'underline'],
            'content': ['added', 'deleted', 'modified', 'moved'],
            'structural': ['table', 'row', 'column', 'section', 'panel'],
            'visual': ['image', 'chart', 'diagram'],
            'metadata': ['author', 'date', 'properties']
        }
        
        # Pre-trained patterns for change classification
        self.patterns = {
            'minor': re.compile(r'(font|color|size|formatting|spacing)', re.IGNORECASE),
            'major': re.compile(r'(added|deleted|removed|new|table|image|section)', re.IGNORECASE),
            'critical': re.compile(r'(TBD|TODO|FIXME|XXX|critical|important|urgent)', re.IGNORECASE)
        }
    
    def classify_change(self, change: Dict) -> Dict:
        """Classify change severity and category"""
        text = f"{change.get('type', '')} {change.get('status', '')} {change.get('change_detail', '')}"
        text_lower = text.lower()
        
        # Calculate severity score
        severity_score = 0
        if self.patterns['critical'].search(text):
            severity_score = 3  # Critical
        elif self.patterns['major'].search(text):
            severity_score = 2  # Major
        elif self.patterns['minor'].search(text):
            severity_score = 1  # Minor
        else:
            severity_score = 0  # Informational
        
        # Determine category
        category = 'other'
        for cat, keywords in self.change_categories.items():
            if any(keyword in text_lower for keyword in keywords):
                category = cat
                break
        
        return {
            'severity': severity_score,
            'category': category,
            'confidence': 0.85  # Placeholder for ML model confidence
        }

class DocumentSimilarityAnalyzer:
    """Advanced document similarity analysis"""
    
    def __init__(self):
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    
    def calculate_similarity(self, doc1_text: str, doc2_text: str) -> Dict:
        """Calculate comprehensive similarity metrics"""
        try:
            # Text similarity using TF-IDF and cosine similarity
            tfidf_matrix = self.vectorizer.fit_transform([doc1_text, doc2_text])
            cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            # Structural similarity
            struct_sim = self._structural_similarity(doc1_text, doc2_text)
            
            # Content preservation score
            content_score = self._content_preservation_score(doc1_text, doc2_text)
            
            return {
                'overall_similarity': round((cosine_sim + struct_sim + content_score) / 3, 3),
                'text_similarity': round(cosine_sim, 3),
                'structural_similarity': round(struct_sim, 3),
                'content_preservation': round(content_score, 3),
                'change_magnitude': round(1 - ((cosine_sim + struct_sim + content_score) / 3), 3)
            }
        except Exception as e:
            print(f"Similarity analysis error: {e}")
            return {'overall_similarity': 0.0, 'error': str(e)}
    
    def _structural_similarity(self, text1: str, text2: str) -> float:
        """Calculate structural similarity based on paragraph and sentence structure"""
        paras1 = text1.split('\n')
        paras2 = text2.split('\n')
        
        if not paras1 or not paras2:
            return 0.0
        
        # Compare paragraph count ratio
        para_ratio = min(len(paras1), len(paras2)) / max(len(paras1), len(paras2))
        
        # Compare average paragraph length similarity
        avg_len1 = sum(len(p) for p in paras1) / len(paras1)
        avg_len2 = sum(len(p) for p in paras2) / len(paras2)
        len_ratio = min(avg_len1, avg_len2) / max(avg_len1, avg_len2) if max(avg_len1, avg_len2) > 0 else 0.0
        
        return (para_ratio + len_ratio) / 2
    
    def _content_preservation_score(self, text1: str, text2: str) -> float:
        """Calculate how much content is preserved vs changed"""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        elif not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0

# Initialize ML components
change_classifier = ChangeClassifier()
similarity_analyzer = DocumentSimilarityAnalyzer()

# ================================
# UTILITY FUNCTIONS
# ================================

def qn(tag: str) -> str:
    """
    Utility function to handle XML namespaces in DOCX documents.
    """
    if tag.startswith('{'):
        return tag
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + tag

def is_encrypted(file_stream) -> bool:
    """
    Check if a document is encrypted without fully loading it.
    """
    try:
        file_stream.seek(0)
        office_file = msoffcrypto.OfficeFile(file_stream)
        return office_file.is_encrypted()
    except Exception:
        return False
    finally:
        file_stream.seek(0)

def load_docx(file_stream, password: Optional[str] = None) -> docx.Document:
    """
    Load DOCX file, handling both encrypted and unencrypted documents.
    """
    file_stream.seek(0)
    
    if password and is_encrypted(file_stream):
        # Handle encrypted document
        decrypted = io.BytesIO()
        office_file = msoffcrypto.OfficeFile(file_stream)
        try:
            office_file.load_key(password=password)
            office_file.decrypt(decrypted)
            decrypted.seek(0)
            return docx.Document(decrypted)
        except Exception as e:
            raise Exception(f"Failed to decrypt document: {str(e)}")
    else:
        # Handle unencrypted document
        try:
            return docx.Document(file_stream)
        except Exception as e:
            if "encrypted" in str(e).lower():
                raise Exception("Document is encrypted but no password was provided")
            raise Exception(f"Failed to load document: {str(e)}")

def update_progress(comparison_id: str, stage: str, progress: int, total: int = 100):
    """Update progress for real-time tracking"""
    if comparison_id not in progress_data:
        progress_data[comparison_id] = {}
    
    progress_data[comparison_id] = {
        'stage': stage,
        'progress': progress,
        'total': total,
        'timestamp': datetime.now().isoformat()
    }

def get_document_fingerprint(doc: docx.Document) -> Dict:
    """Create a comprehensive fingerprint of the document"""
    content = extract_docx_content_enhanced(doc)
    images = get_docx_images(doc)
    
    # Create fingerprint based on content and structure
    text_content = " ".join([item['text'] for item in content if 'text' in item])
    fingerprint = {
        'total_paragraphs': len([item for item in content if item['type'] == 'paragraph']),
        'total_tables': len([item for item in content if item['type'] == 'table']),
        'total_images': len(images),
        'content_hash': md5(text_content.encode()).hexdigest(),
        'structure_hash': md5(str([(item['type'], item.get('index', 0)) for item in content]).encode()).hexdigest()
    }
    
    return fingerprint

def extract_full_text(doc: docx.Document) -> str:
    """Extract all text content from document for similarity analysis"""
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return "\n".join(full_text)

# ================================
# CONTENT EXTRACTION FUNCTIONS
# ================================

def extract_docx_content_enhanced(doc: docx.Document) -> List[Dict[str, Any]]:
    """
    Extract content from DOCX with enhanced strikethrough text handling.
    """
    content = []
    
    # Extract paragraphs with enhanced formatting detection
    for para_idx, para in enumerate(doc.paragraphs, start=1):
        if para.text.strip():
            runs_data = []
            strikethrough_found = False
            strikethrough_texts = []
            
            # Analyze each text run in the paragraph
            for run_idx, run in enumerate(para.runs):
                run_text = run.text.strip()
                if run_text:
                    # Detect formatting properties
                    strike = detect_strikethrough(run)
                    
                    run_data = {
                        "text": run_text,
                        "strikethrough": strike,
                        "bold": run.font.bold,
                        "italic": run.font.italic,
                        "underline": run.font.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size,
                        "font_color": run.font.color.rgb if run.font.color else None,
                    }
                    
                    if strike:
                        strikethrough_found = True
                        strikethrough_texts.append(run_text)
                    
                    runs_data.append(run_data)
            
            # Combine strikethrough texts
            strikethrough_text_combined = " | ".join(strikethrough_texts) if strikethrough_texts else ""
            
            content.append({
                "type": "paragraph",
                "index": para_idx,
                "text": para.text,
                "runs": runs_data,
                "style": para.style.name if para.style else "Normal",
                "has_strikethrough": strikethrough_found,
                "strikethrough_texts": strikethrough_texts,
                "strikethrough_combined": strikethrough_text_combined
            })
    
    # Extract tables (excluding first two tables which are often templates)
    for table_idx, table in enumerate(doc.tables, start=1):
        if table_idx <= 2:  # Skip first two tables
            continue
            
        table_data = {
            "type": "table",
            "index": table_idx,
            "rows": [],
            "has_strikethrough": False,
            "strikethrough_texts": []
        }
        
        # Process each row and cell in the table
        for row_idx, row in enumerate(table.rows, start=1):
            row_data = {"cells": []}
            for cell_idx, cell in enumerate(row.cells, start=1):
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                
                cell_has_strikethrough = False
                cell_strikethrough_texts = []
                
                # Check for strikethrough in cell content
                for para in cell.paragraphs:
                    for run in para.runs:
                        run_text = run.text.strip()
                        if run_text:
                            strike = detect_strikethrough(run)
                            if strike:
                                cell_has_strikethrough = True
                                cell_strikethrough_texts.append(run_text)
                                table_data["has_strikethrough"] = True
                                table_data["strikethrough_texts"].append(run_text)
                
                # Combine strikethrough texts
                cell_strikethrough_combined = " | ".join(cell_strikethrough_texts) if cell_strikethrough_texts else ""
                
                row_data["cells"].append({
                    "text": cell_text,
                    "row": row_idx,
                    "col": cell_idx,
                    "contains_tbd": "TBD" in cell_text.upper(),
                    "contains_code_names": bool(re.search(r'\[[A-Z]+\]', cell_text)),
                    "urls": re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', cell_text),
                    "has_strikethrough": cell_has_strikethrough,
                    "strikethrough_texts": cell_strikethrough_texts,
                    "strikethrough_combined": cell_strikethrough_combined
                })
            
            table_data["rows"].append(row_data)
        
        content.append(table_data)
    
    return content

def detect_strikethrough(run) -> bool:
    """
    Detect if text has strikethrough formatting using multiple methods.
    """
    try:
        # Method 1: Check font property directly
        if run.font.strike:
            return True
            
        # Method 2: Check XML directly for more reliable detection
        if hasattr(run, '_element'):
            # Check for w:strike elements
            strike_elements = run._element.xpath('.//w:strike')
            if strike_elements:
                for strike_elm in strike_elements:
                    val = strike_elm.get(qn('w:val'))
                    if val is None or val != "false":
                        return True
            
            # Check for w:dstrike elements (double strikethrough)
            dstrike_elements = run._element.xpath('.//w:dstrike')
            if dstrike_elements:
                for dstrike_elem in dstrike_elements:
                    val = dstrike_elem.get(qn('w:val'))
                    if val is None or val in ['true', '1', 'on']:
                        return True
        
        # Method 3: Check run properties via XML
        if hasattr(run, '_r'):
            r_pr = run._r.rPr
            if r_pr is not None:
                strike = r_pr.find(qn('w:strike'))
                if strike is not None:
                    val = strike.get(qn('w:val'))
                    if val is None or val not in ['false', '0', 'off']:
                        return True
                    
        return False
    except Exception as e:
        print(f"Error in detect_strikethrough: {e}")
        return False

def detect_panels(doc: docx.Document) -> List[Dict[str, Any]]:
    """
    Detect panels in document based on section breaks, headings, or other markers.
    """
    panels = []
    
    # Look for section breaks
    for sect in doc.sections:
        panels.append({
            "type": "section",
            "start": "beginning",
            "content": "Section content"
        })
    
    # Look for headings that might indicate panels
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            panels.append({
                "type": "heading",
                "level": int(para.style.name.replace('Heading', '')) if para.style.name.replace('Heading', '').isdigit() else 1,
                "text": para.text,
                "position": len(panels) + 1
            })
    
    return panels

# ================================
# IMAGE HANDLING FUNCTIONS
# ================================

def get_docx_images(doc: docx.Document) -> List[Dict[str, Any]]:
    """
    Extract all images from DOCX document with comprehensive detection.
    """
    images = []
    
    try:
        # Method 1: Check document relationships
        if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        img_bytes = rel.target_part.blob
                        img_hash = md5(img_bytes).hexdigest()
                        
                        # Get image dimensions using PIL
                        pil_img = PILImage.open(io.BytesIO(img_bytes))
                        width, height = pil_img.size
                        
                        images.append({
                            "hash": img_hash,
                            "bytes": img_bytes,
                            "width": width,
                            "height": height,
                            "format": pil_img.format,
                            "rel_id": rel_id,
                            "source": "document_relationships"
                        })
                    except Exception as e:
                        continue
        
        # Method 2: Check inline shapes
        try:
            for element in doc.element.body:
                for inline in element.xpath('.//wp:inline', namespaces={
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                }):
                    blip = inline.xpath('.//a:blip', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    })
                    if blip:
                        embed_id = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in doc.part.rels:
                            try:
                                img_part = doc.part.rels[embed_id].target_part
                                img_bytes = img_part.blob
                                img_hash = md5(img_bytes).hexdigest()
                                
                                pil_img = PILImage.open(io.BytesIO(img_bytes))
                                width, height = pil_img.size
                                
                                images.append({
                                    "hash": img_hash,
                                    "bytes": img_bytes,
                                    "width": width,
                                    "height": height,
                                    "format": pil_img.format,
                                    "embed_id": embed_id,
                                    "source": "inline_shapes"
                                })
                            except Exception as e:
                                continue
        except Exception as e:
            pass
        
    except Exception as e:
        print(f"Error in get_docx_images: {e}")
    
    return images

def get_table_images(doc: docx.Document) -> List[Dict[str, Any]]:
    """
    Extract images from inside table cells with proper XML parsing.
    """
    table_images = []
    
    try:
        # Namespace map for XML parsing
        nsmap = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        for table_idx, table in enumerate(doc.tables, start=1):
            if table_idx <= 2:  # Skip first two tables
                continue
            
            for row_idx, row in enumerate(table.rows, start=1):
                for cell_idx, cell in enumerate(row.cells, start=1):
                    # Convert cell to XML element for parsing
                    cell_xml = cell._element
                    
                    # Find all graphic elements in the cell
                    for graphic in cell_xml.xpath('.//w:drawing//a:graphic', namespaces=nsmap):
                        # Find blip elements which reference images
                        for blip in graphic.xpath('.//a:blip', namespaces=nsmap):
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id and embed_id in doc.part.rels:
                                try:
                                    img_part = doc.part.rels[embed_id].target_part
                                    img_bytes = img_part.blob
                                    img_hash = md5(img_bytes).hexdigest()
                                    
                                    # Get image dimensions
                                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                                    width, height = pil_img.size
                                    
                                    table_images.append({
                                        "table": table_idx,
                                        "row": row_idx,
                                        "col": cell_idx,
                                        "hash": img_hash,
                                        "bytes": img_bytes,
                                        "width": width,
                                        "height": height,
                                        "format": pil_img.format,
                                        "embed_id": embed_id
                                    })
                                except Exception as e:
                                    continue
    except Exception as e:
        print(f"Error extracting table images: {e}")
    
    return table_images

def insert_image_into_excel(ws, img_bytes: bytes, cell_address: str, 
                           max_width: int = 120, max_height: int = 120) -> bool:
    """
    Insert an image into Excel worksheet with size adjustment.
    """
    if not img_bytes:
        return False
        
    try:
        # Create temporary file for the image
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
            # Open and process the image
            pil_img = PILImage.open(io.BytesIO(img_bytes))
            
            # Convert to RGB if necessary
            if pil_img.mode in ('RGBA', 'LA', 'P'):
                background = PILImage.new('RGB', pil_img.size, (255, 255, 255))
                if pil_img.mode == 'P':
                    pil_img = pil_img.convert('RGBA')
                background.paste(pil_img, mask=pil_img.split()[-1] if pil_img.mode == 'RGBA' else None)
                pil_img = background
            
            # Resize image if too large
            original_width, original_height = pil_img.size
            width_ratio = max_width / original_width
            height_ratio = max_height / original_height
            ratio = min(width_ratio, height_ratio, 1.0)  # Don't enlarge small images
            
            if ratio < 1.0:
                new_width = int(original_width * ratio)
                new_height = int(original_height * ratio)
                pil_img = pil_img.resize((new_width, new_height), PILImage.Resampling.LANCZOS)
            
            # Save as PNG for better compatibility
            pil_img.save(tmp_file.name, format="PNG", optimize=True)
            
            # Create OpenPyXL image and add to worksheet
            img = XLImage(tmp_file.name)
            img.anchor = cell_address
            
            # Adjust cell size to fit image
            col_letter = cell_address[0]  # Get column letter
            row_num = int(cell_address[1:])  # Get row number
            
            # Set column width based on image width
            col_width = max(8.43, min(50, img.width * 0.14))
            ws.column_dimensions[col_letter].width = col_width
            
            # Set row height based on image height
            row_height = max(15, min(400, img.height * 0.75))
            ws.row_dimensions[row_num].height = row_height
            
            ws.add_image(img)
            
            # Clean up temporary file
            os.unlink(tmp_file.name)
            return True
            
    except Exception as e:
        print(f"Error inserting image into Excel: {e}")
        if 'tmp_file' in locals() and os.path.exists(tmp_file.name):
            os.unlink(tmp_file.name)
        return False

# ================================
# COMPARISON FUNCTIONS
# ================================

def compare_images_enhanced(imgs1: List[Dict], imgs2: List[Dict]) -> List[Dict[str, Any]]:
    """
    Enhanced image comparison that properly captures deleted images.
    """
    changes = []
    
    # Group images by hash for comparison
    hashes1 = {img["hash"]: img for img in imgs1}
    hashes2 = {img["hash"]: img for img in imgs2}
    
    # Find deleted images (in doc1 but not in doc2)
    for img_hash, img1 in hashes1.items():
        if img_hash not in hashes2:
            changes.append({
                "type": "image",
                "status": "deleted",
                "hash": img_hash,
                "old_img": img1["bytes"],
                "new_img": None,
                "width": img1["width"],
                "height": img1["height"],
                "alignment": "document_level",
                "change_detail": f"Image deleted - Original size: {img1['width']}x{img1['height']}"
            })
    
    # Find added images (in doc2 but not in doc1)
    for img_hash, img2 in hashes2.items():
        if img_hash not in hashes1:
            changes.append({
                "type": "image",
                "status": "added",
                "hash": img_hash,
                "old_img": None,
                "new_img": img2["bytes"],
                "width": img2["width"],
                "height": img2["height"],
                "alignment": "document_level",
                "change_detail": f"Image added - Size: {img2['width']}x{img2['height']}"
            })
    
    # Find modified images (same hash but different properties)
    common_hashes = set(hashes1.keys()) & set(hashes2.keys())
    for img_hash in common_hashes:
        img1 = hashes1[img_hash]
        img2 = hashes2[img_hash]
        
        # Check for size changes
        if img1["width"] != img2["width"] or img1["height"] != img2["height"]:
            changes.append({
                "type": "image",
                "status": "modified",
                "hash": img_hash,
                "old_img": img1["bytes"],
                "new_img": img2["bytes"],
                "width_original": img1["width"],
                "height_original": img1["height"],
                "width_modified": img2["width"],
                "height_modified": img2["height"],
                "alignment": "document_level",
                "change_detail": f"Size changed from {img1['width']}x{img1['height']} to {img2['width']}x{img2['height']}"
            })
    
    return changes

def compare_table_images_enhanced(imgs1: List[Dict], imgs2: List[Dict]) -> List[Dict[str, Any]]:
    """
    Enhanced table image comparison that properly captures deleted images.
    """
    changes = []
    
    # Group by table, row, column AND hash
    all_images1 = {(img["table"], img["row"], img["col"], img["hash"]): img for img in imgs1}
    all_images2 = {(img["table"], img["row"], img["col"], img["hash"]): img for img in imgs2}
    
    # Find deleted table images
    for img_key, img1 in all_images1.items():
        if img_key not in all_images2:
            table, row, col, img_hash = img_key
            changes.append({
                "type": "table_image",
                "status": "deleted",
                "table": table,
                "row": row,
                "col": col,
                "hash": img_hash,
                "old_img": img1["bytes"],
                "new_img": None,
                "width": img1["width"],
                "height": img1["height"],
                "alignment": f"Table {table}, Cell ({row},{col})",
                "change_detail": f"Table image deleted from Table {table}, Cell ({row},{col})"
            })
    
    # Find added table images
    for img_key, img2 in all_images2.items():
        if img_key not in all_images1:
            table, row, col, img_hash = img_key
            changes.append({
                "type": "table_image",
                "status": "added",
                "table": table,
                "row": row,
                "col": col,
                "hash": img_hash,
                "old_img": None,
                "new_img": img2["bytes"],
                "width": img2["width"],
                "height": img2["height"],
                "alignment": f"Table {table}, Cell ({row},{col})",
                "change_detail": f"Table image added to Table {table}, Cell ({row},{col})"
            })
    
    return changes

def compare_text_content_enhanced(content1: List[Dict], content2: List[Dict]) -> List[Dict[str, Any]]:
    """
    Enhanced text comparison with detailed change tracking and strikethrough detection.
    """
    changes = []
    strikethrough_changes = []
    processed_strikethrough_paragraphs = set()
    
    # Compare paragraphs
    para1 = [item for item in content1 if item["type"] == "paragraph"]
    para2 = [item for item in content2 if item["type"] == "paragraph"]
    
    for i in range(max(len(para1), len(para2))):
        if i < len(para1) and i < len(para2):
            p1 = para1[i]
            p2 = para2[i]
            
            # Get strikethrough information
            strike_original = p1["has_strikethrough"]
            strike_modified = p2["has_strikethrough"]
            strike_text_original = p1["strikethrough_combined"]
            strike_text_modified = p2["strikethrough_combined"]
            
            # Enhanced strikethrough detection
            strikethrough_runs_original = [run for run in p1["runs"] if run.get("strikethrough", False)]
            strikethrough_runs_modified = [run for run in p2["runs"] if run.get("strikethrough", False)]
            
            strikethrough_original = len(strikethrough_runs_original) > 0
            strikethrough_modified = len(strikethrough_runs_modified) > 0
            
            # Capture strikethrough text content
            strikethrough_text_original = " | ".join([run["text"] for run in strikethrough_runs_original])
            strikethrough_text_modified = " | ".join([run["text"] for run in strikethrough_runs_modified])
            
            paragraph_key = f"para_{i+1}"
            
            # SEPARATE STRIKETHROUGH CHANGES
            if (strike_original or strike_modified) and paragraph_key not in processed_strikethrough_paragraphs:
                if strike_original and strike_text_original.strip():
                    strikethrough_changes.append({
                        "type": "paragraph_strikethrough",
                        "index": i + 1,
                        "status": "strikethrough_original",
                        "original": p1["text"],
                        "modified": p2["text"],
                        "strikethrough_original": True,
                        "strikethrough_text_original": strike_text_original,
                        "strikethrough_modified": strike_modified,
                        "strikethrough_text_modified": strike_text_modified,
                        "change_detail": f"Strikethrough text found in original: {strike_text_original}",
                    })
                    processed_strikethrough_paragraphs.add(paragraph_key)
                
                if strike_modified and strike_text_modified.strip():
                    if not strike_original or strike_text_original != strike_text_modified:
                        strikethrough_changes.append({
                            "type": "paragraph_strikethrough",
                            "index": i + 1,
                            "status": "strikethrough_modified",
                            "original": p1["text"],
                            "modified": p2["text"],
                            "strikethrough_original": strike_original,
                            "strikethrough_text_original": strike_text_original,
                            "strikethrough_modified": True,
                            "strikethrough_text_modified": strike_text_modified,
                            "change_detail": f"Strikethrough text found in modified: {strike_text_modified}",
                        })
                        
            # Check for text changes
            if p1["text"] != p2["text"]:
                diff = list(difflib.ndiff(p1["text"].split(), p2["text"].split()))
                added = [word[2:] for word in diff if word.startswith('+ ')]
                removed = [word[2:] for word in diff if word.startswith('- ')]
                
                # Check for various content patterns
                double_spaces_original = len(re.findall(r'  ', p1["text"]))
                double_spaces_modified = len(re.findall(r'  ', p2["text"]))
                
                tbd_original = "TBD" in p1["text"].upper()
                tbd_modified = "TBD" in p2["text"].upper()
                
                code_names_original = bool(re.search(r'\[[A-Z]+\]', p1["text"]))
                code_names_modified = bool(re.search(r'\[[A-Z]+\]', p2["text"]))
                
                urls_original = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', p1["text"])
                urls_modified = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', p2["text"])
                
                has_non_strikethrough_changes = (
                    added or removed or 
                    double_spaces_original != double_spaces_modified or
                    tbd_original != tbd_modified or
                    code_names_original != code_names_modified or
                    urls_original != urls_modified
                )
                
                if has_non_strikethrough_changes:
                    changes.append({
                        "type": "paragraph",
                        "index": i + 1,
                        "status": "modified",
                        "original": p1["text"],
                        "modified": p2["text"],
                        "added": " ".join(added),
                        "removed": " ".join(removed),
                        "double_spaces_original": double_spaces_original,
                        "double_spaces_modified": double_spaces_modified,
                        "tbd_original": tbd_original,
                        "tbd_modified": tbd_modified,
                        "code_names_original": code_names_original,
                        "code_names_modified": code_names_modified,
                        "urls_original": urls_original,
                        "urls_modified": urls_modified,
                    })
            
            # Check for formatting changes even if text is the same
            elif p1["text"] == p2["text"] and paragraph_key not in processed_strikethrough_paragraphs:
                formatting_changes = []
                strikethrough_format_changes = []
                
                for run_idx in range(max(len(p1["runs"]), len(p2["runs"]))):
                    if run_idx < len(p1["runs"]) and run_idx < len(p2["runs"]):
                        run1 = p1["runs"][run_idx]
                        run2 = p2["runs"][run_idx]
                        
                        if run1["text"] == run2["text"]:
                            format_diffs = []
                            non_strikethrough_diffs = []
                            
                            if run1.get("bold") != run2.get("bold"):
                                format_diffs.append("bold")
                                non_strikethrough_diffs.append("bold")
                            if run1.get("italic") != run2.get("italic"):
                                format_diffs.append("italic")
                                non_strikethrough_diffs.append("italic")
                            if run1.get("underline") != run2.get("underline"):
                                format_diffs.append("underline")
                                non_strikethrough_diffs.append("underline")
                            
                            # Handle strikethrough formatting separately
                            if run1.get("strikethrough") != run2.get("strikethrough"):
                                format_diffs.append("strikethrough")
                                if run1.get("strikethrough") and not run2.get("strikethrough"):
                                    strikethrough_format_changes.append(f"strikethrough removed: '{run1['text']}'")
                                elif not run1.get("strikethrough") and run2.get("strikethrough"):
                                    strikethrough_format_changes.append(f"Strikethrough added: '{run1['text']}'")
                            
                            if run1.get("font_name") != run2.get("font_name"):
                                format_diffs.append("font")
                                non_strikethrough_diffs.append("font")
                            if run1.get("font_size") != run2.get("font_size"):
                                format_diffs.append("font size")
                                non_strikethrough_diffs.append("font size")
                            if run1.get("font_color") != run2.get("font_color"):
                                format_diffs.append("font color")
                                non_strikethrough_diffs.append("font color")
                                
                            if format_diffs:
                                if non_strikethrough_diffs:
                                    formatting_changes.append({
                                        "text": run1["text"],
                                        "changes": non_strikethrough_diffs
                                    })
                
                # Add strikethrough formatting changes to strikethrough_changes list
                if strikethrough_format_changes and paragraph_key not in processed_strikethrough_paragraphs:
                    unique_strikethrough_changes = list(set(strikethrough_format_changes))
                    strikethrough_changes.append({
                        "type": "paragraph_formatting_strikethrough",
                        "index": i + 1,
                        "status": "strikethrough_formatting_changed",
                        "original": p1["text"],
                        "modified": p2["text"],
                        "formatting_changes": unique_strikethrough_changes,
                        "change_detail": "Strikethrough formatting changes:\n" + "\n".join(unique_strikethrough_changes),
                    })
                    processed_strikethrough_paragraphs.add(paragraph_key)
                
                # Only add to regular changes if there are non-strikethrough formatting changes
                if formatting_changes:
                    change_detail = "Formatting changes:\n"
                    for fmt_change in formatting_changes:
                        change_detail += f"'{fmt_change['text']}': {', '.join(fmt_change['changes'])}\n"
                    
                    changes.append({
                        "type": "paragraph_formatting",
                        "index": i + 1,
                        "status": "formatting_changed",
                        "original": p1["text"],
                        "modified": p2["text"],
                        "formatting_changes": formatting_changes,
                        "change_detail": change_detail,
                    })
                    
        elif i < len(para1):  # Deleted paragraph
            p1 = para1[i]
            # Check for strikethrough in deleted paragraph
            strikethrough_runs = [run for run in p1["runs"] if run.get("strikethrough", False)]
            strikethrough_text = " | ".join([run["text"] for run in strikethrough_runs])
            
            # Add strikethrough info to strikethrough_changes if present
            if strikethrough_text.strip():
                strikethrough_changes.append({
                    "type": "paragraph_strikethrough",
                    "index": i + 1,
                    "status": "deleted_strikethrough",
                    "original": p1["text"],
                    "modified": "",
                    "strikethrough_original": True,
                    "strikethrough_text_original": strikethrough_text,
                    "change_detail": f"Deleted paragraph contained strikethrough: {strikethrough_text}",
                })
            
            # Always add to regular changes for deleted paragraphs
            changes.append({
                "type": "paragraph",
                "index": i + 1,
                "status": "deleted",
                "original": p1["text"],
                "modified": "",
                "added": "",
                "removed": p1["text"],
            })
                
        elif i < len(para2):  # Added paragraph
            p2 = para2[i]
            # Check for strikethrough in added paragraph
            strikethrough_runs = [run for run in p2["runs"] if run.get("strikethrough", False)]
            strikethrough_text = " | ".join([run["text"] for run in strikethrough_runs])
            
            # Add strikethrough info to strikethrough_changes if present
            if strikethrough_text.strip():
                strikethrough_changes.append({
                    "type": "paragraph_strikethrough",
                    "index": i + 1,
                    "status": "added_strikethrough",
                    "original": "",
                    "modified": p2["text"],
                    "strikethrough_modified": True,
                    "strikethrough_text_modified": strikethrough_text,
                    "change_detail": f"Added paragraph contains strikethrough: {strikethrough_text}",
                })
            
            # Always add to regular changes for added paragraphs
            changes.append({
                "type": "paragraph",
                "index": i + 1,
                "status": "added",
                "original": "",
                "modified": p2["text"],
                "added": p2["text"],
                "removed": "",
            })
    
    # Remove duplicate strikethrough changes by paragraph
    unique_strikethrough_changes = []
    seen_paragraphs = set()
    
    for change in strikethrough_changes:
        para_key = f"para_{change.get('index')}"
        if para_key not in seen_paragraphs:
            unique_strikethrough_changes.append(change)
            seen_paragraphs.add(para_key)
    
    # Combine regular changes with strikethrough changes
    all_changes = changes + unique_strikethrough_changes
    return all_changes

def compare_tables(content1: List[Dict], content2: List[Dict]) -> List[Dict[str, Any]]:
    """
    Compare tables between two documents without duplication.
    """
    changes = []
    strikethrough_changes = []
    
    tables1 = [item for item in content1 if item["type"] == "table"]
    tables2 = [item for item in content2 if item["type"] == "table"]
    
    for i in range(max(len(tables1), len(tables2))):
        if i < len(tables1) and i < len(tables2):
            table1 = tables1[i]
            table2 = tables2[i]
            
            # Compare table-level strikethrough
            if table1["has_strikethrough"] or table2["has_strikethrough"]:
                strike1_text = " | ".join(table1.get("strikethrough_texts", []))
                strike2_text = " | ".join(table2.get("strikethrough_texts", []))
                
                if strike1_text != strike2_text:
                    strikethrough_changes.append({
                        "type": "table_strikethrough",
                        "table_index": i + 1,
                        "status": "modified",
                        "original_strikethrough": strike1_text,
                        "modified_strikethrough": strike2_text,
                        "original_has_strike": table1["has_strikethrough"],
                        "modified_has_strike": table2["has_strikethrough"],
                        "change_detail": f"Table strikethrough content changed"
                    })

            # Compare row count
            if len(table1["rows"]) != len(table2["rows"]):
                changes.append({
                    "type": "table",
                    "index": i + 1,
                    "status": "structure_changed",
                    "detail": f"Row count changed from {len(table1['rows'])} to {len(table2['rows'])}",
                    "original": f"Table with {len(table1['rows'])} rows",
                    "modified": f"Table with {len(table2['rows'])} rows"
                })
                
            # Compare cell content
            for r_idx in range(max(len(table1["rows"]), len(table2["rows"]))):
                if r_idx < len(table1["rows"]) and r_idx < len(table2["rows"]):
                    row1 = table1["rows"][r_idx]
                    row2 = table2["rows"][r_idx]
                    
                    # Compare column count
                    if len(row1["cells"]) != len(row2["cells"]):
                        changes.append({
                            "type": "table",
                            "index": i + 1,
                            "status": "structure_changed",
                            "detail": f"Column count changed in row {r_idx+1} from {len(row1['cells'])} to {len(row2['cells'])}",
                            "original": f"Row with {len(row1['cells'])} columns",
                            "modified": f"Row with {len(row2['cells'])} columns"
                        })
                    
                    # Compare cell content
                    for c_idx in range(max(len(row1["cells"]), len(row2["cells"]))):
                        if c_idx < len(row1["cells"]) and c_idx < len(row2["cells"]):
                            cell1 = row1["cells"][c_idx]
                            cell2 = row2["cells"][c_idx]
                            
                            # Check for strikethrough changes separately
                            strike1 = " | ".join(cell1.get("strikethrough_texts", []))
                            strike2 = " | ".join(cell2.get("strikethrough_texts", []))
                            
                            has_strikethrough_changes = (
                                cell1["has_strikethrough"] != cell2["has_strikethrough"] or
                                (cell1["has_strikethrough"] and cell2["has_strikethrough"] and strike1 != strike2)
                            )
                            
                            has_text_changes = cell1["text"] != cell2["text"]
                            
                            # Handle strikethrough changes separately
                            if has_strikethrough_changes:
                                strikethrough_changes.append({
                                    "type": "table_cell_strikethrough",
                                    "table_index": i + 1,
                                    "row": r_idx + 1,
                                    "col": c_idx + 1,
                                    "status": "strikethrough_modified",
                                    "original": cell1["text"],
                                    "modified": cell2["text"],
                                    "strikethrough_original": strike1,
                                    "strikethrough_modified": strike2,
                                    "has_strikethrough_original": cell1["has_strikethrough"],
                                    "has_strikethrough_modified": cell2["has_strikethrough"],
                                    "change_detail": "Strikethrough content changed in table cell"
                                })
                            
                            # Handle regular text changes
                            if has_text_changes and not has_strikethrough_changes:
                                changes.append({
                                    "type": "table_cell",
                                    "table_index": i + 1,
                                    "row": r_idx + 1,
                                    "col": c_idx + 1,
                                    "status": "modified",
                                    "original": cell1["text"],
                                    "modified": cell2["text"],
                                    "change_detail": "Text content modified in table cell"
                                })
                            elif has_text_changes and has_strikethrough_changes:
                                changes.append({
                                    "type": "table_cell",
                                    "table_index": i + 1,
                                    "row": r_idx + 1,
                                    "col": c_idx + 1,
                                    "status": "modified",
                                    "original": cell1["text"],
                                    "modified": cell2["text"],
                                    "change_detail": "Text content modified in table cell"
                                })
                                
                        elif c_idx < len(row1["cells"]):
                            # Cell deleted
                            cell1 = row1["cells"][c_idx]
                            strike1 = " | ".join(cell1.get("strikethrough_texts", []))
                            
                            if cell1["has_strikethrough"] and strike1.strip():
                                strikethrough_changes.append({
                                    "type": "table_cell_strikethrough",
                                    "table_index": i + 1,
                                    "row": r_idx + 1,
                                    "col": c_idx + 1,
                                    "status": "deleted_strikethrough",
                                    "original": cell1["text"],
                                    "modified": "",
                                    "strikethrough_original": strike1,
                                    "has_strikethrough_original": True,
                                    "change_detail": "Deleted cell contained strikethrough text"
                                })
                            
                            changes.append({
                                "type": "table_cell",
                                "table_index": i + 1,
                                "row": r_idx + 1,
                                "col": c_idx + 1,
                                "status": "deleted",
                                "original": cell1["text"],
                                "modified": "",
                                "change_detail": "Cell deleted from table"
                            })
                        elif c_idx < len(row2["cells"]):
                            # Cell added
                            cell2 = row2["cells"][c_idx]
                            strike2 = " | ".join(cell2.get("strikethrough_texts", []))
                            
                            if cell2["has_strikethrough"] and strike2.strip():
                                strikethrough_changes.append({
                                    "type": "table_cell_strikethrough",
                                    "table_index": i + 1,
                                    "row": r_idx + 1,
                                    "col": c_idx + 1,
                                    "status": "added_strikethrough",
                                    "original": "",
                                    "modified": cell2["text"],
                                    "strikethrough_modified": strike2,
                                    "has_strikethrough_modified": True,
                                    "change_detail": "Added cell contains strikethrough text"
                                })
                            
                            changes.append({
                                "type": "table_cell",
                                "table_index": i + 1,
                                "row": r_idx + 1,
                                "col": c_idx + 1,
                                "status": "added",
                                "original": "",
                                "modified": cell2["text"],
                                "change_detail": "Cell added to table"
                            })
                            
                elif r_idx < len(table1["rows"]):
                    # Entire row deleted
                    row1 = table1["rows"][r_idx]
                    changes.append({
                        "type": "table_row",
                        "table_index": i + 1,
                        "row": r_idx + 1,
                        "status": "deleted",
                        "original": f"Row {r_idx+1} with {len(row1['cells'])} columns",
                        "modified": "",
                        "change_detail": f"Row {r_idx+1} deleted from table"
                    })
                elif r_idx < len(table2["rows"]):
                    # Entire row added
                    row2 = table2["rows"][r_idx]
                    changes.append({
                        "type": "table_row",
                        "table_index": i + 1,
                        "row": r_idx + 1,
                        "status": "added",
                        "original": "",
                        "modified": f"Row {r_idx+1} with {len(row2['cells'])} columns",
                        "change_detail": f"Row {r_idx+1} added to table"
                    })
                    
        elif i < len(tables1):
            # Entire table deleted
            changes.append({
                "type": "table",
                "index": i + 1,
                "status": "deleted",
                "original": f"Table with {len(tables1[i]['rows'])} rows",
                "modified": "",
                "change_detail": f"Table {i+1} deleted from document"
            })
        elif i < len(tables2):
            # Entire table added
            changes.append({
                "type": "table",
                "index": i + 1,
                "status": "added",
                "original": "",
                "modified": f"Table with {len(tables2[i]['rows'])} rows",
                "change_detail": f"Table {i+1} added to document"
            })
    
    # Combine regular changes with strikethrough changes
    all_changes = changes + strikethrough_changes
    return all_changes

def compare_panels(panels1: List[Dict], panels2: List[Dict]) -> List[Dict[str, Any]]:
    """
    Compare panels between documents.
    """
    changes = []
    
    # Simple comparison based on panel count and content
    if len(panels1) != len(panels2):
        changes.append({
            "type": "panel_count",
            "status": "changed",
            "original_count": len(panels1),
            "modified_count": len(panels2)
        })
    
    # Compare panel content
    for i in range(max(len(panels1), len(panels2))):
        if i < len(panels1) and i < len(panels2):
            p1 = panels1[i]
            p2 = panels2[i]
            
            if p1.get("text") != p2.get("text"):
                changes.append({
                    "type": "panel_content",
                    "index": i + 1,
                    "status": "modified",
                    "original": p1.get("text", ""),
                    "modified": p2.get("text", "")
                })
                
            if p1.get("type") != p2.get("type"):
                changes.append({
                    "type": "panel_type",
                    "index": i + 1,
                    "status": "modified",
                    "original": p1.get("type", ""),
                    "modified": p2.get("type", "")
                })
                
        elif i < len(panels1):
            changes.append({
                "type": "panel",
                "index": i + 1,
                "status": "deleted",
                "original": panels1[i].get("text", ""),
                "modified": ""
            })
        elif i < len(panels2):
            changes.append({
                "type": "panel",
                "index": i + 1,
                "status": "added",
                "original": "",
                "modified": panels2[i].get("text", "")
            })
    
    return changes

def compare_docx_enhanced(doc1: docx.Document, doc2: docx.Document, comparison_id: str = None) -> Tuple[List[Dict], Dict[str, Any]]:
    """
    Enhanced DOCX comparison with ML-powered analysis and progress tracking.
    """
    if comparison_id:
        update_progress(comparison_id, "Extracting content", 10)
    
    # Extract content with enhanced analysis
    content1 = extract_docx_content_enhanced(doc1)
    content2 = extract_docx_content_enhanced(doc2)
    
    if comparison_id:
        update_progress(comparison_id, "Analyzing images", 30)
    
    # Extract images
    imgs1 = get_docx_images(doc1)
    imgs2 = get_docx_images(doc2)
    table_imgs1 = get_table_images(doc1)
    table_imgs2 = get_table_images(doc2)
    
    if comparison_id:
        update_progress(comparison_id, "Detecting panels", 40)
    
    # Detect panels
    panels1 = detect_panels(doc1)
    panels2 = detect_panels(doc2)
    
    if comparison_id:
        update_progress(comparison_id, "Comparing text", 50)
    
    # Compare different aspects
    text_changes = compare_text_content_enhanced(content1, content2)
    
    if comparison_id:
        update_progress(comparison_id, "Comparing tables", 60)
    
    table_changes = compare_tables(content1, content2)
    
    if comparison_id:
        update_progress(comparison_id, "Comparing images", 70)
    
    image_changes = compare_images_enhanced(imgs1, imgs2)
    table_image_changes = compare_table_images_enhanced(table_imgs1, table_imgs2)
    
    if comparison_id:
        update_progress(comparison_id, "Comparing panels", 80)
    
    panel_changes = compare_panels(panels1, panels2)
    
    # Combine all changes
    all_changes = text_changes + table_changes + image_changes + table_image_changes + panel_changes
    all_changes = remove_duplicate_changes(all_changes)
    
    if comparison_id:
        update_progress(comparison_id, "Analyzing changes", 90)
    
    # Enhanced change analysis with ML
    analyzed_changes = []
    for change in all_changes:
        classification = change_classifier.classify_change(change)
        change.update({
            'severity': classification['severity'],
            'category': classification['category'],
            'confidence': classification['confidence']
        })
        analyzed_changes.append(change)
    
    # Calculate document similarity
    doc1_text = extract_full_text(doc1)
    doc2_text = extract_full_text(doc2)
    similarity_metrics = similarity_analyzer.calculate_similarity(doc1_text, doc2_text)
    
    # Enhanced statistics
    para_count = len([item for item in content1 if item["type"] == "paragraph"])
    processed_tables1 = [item for item in content1 if item["type"] == "table"]
    table_count = len(processed_tables1)
    image_count = len(imgs1)
    table_image_count = len(table_imgs1)
    panel_count = len(panels1)
    
    # Enhanced change analysis
    change_counts = Counter([change['type'] for change in analyzed_changes])
    severity_counts = Counter([change['severity'] for change in analyzed_changes])
    category_counts = Counter([change['category'] for change in analyzed_changes])
    
    # Calculate change density
    total_elements = para_count + table_count + panel_count
    change_density = len(analyzed_changes) / total_elements if total_elements > 0 else 0
    
    stats = {
        "paragraphs": para_count,
        "tables": table_count,
        "images": image_count,
        "table_images": table_image_count,
        "panels": panel_count,
        "total_elements": total_elements,
        "changes": {
            "total": len(analyzed_changes),
            "by_type": dict(change_counts),
            "by_severity": dict(severity_counts),
            "by_category": dict(category_counts)
        },
        "change_density": round(change_density, 3),
        "similarity_metrics": similarity_metrics,
        "fingerprint_doc1": get_document_fingerprint(doc1),
        "fingerprint_doc2": get_document_fingerprint(doc2),
        "processing_time": datetime.now().isoformat()
    }
    
    if comparison_id:
        update_progress(comparison_id, "Complete", 100)
    
    return analyzed_changes, stats

def remove_duplicate_changes(changes: List[Dict]) -> List[Dict]:
    """
    Remove duplicate changes based on unique identifiers.
    """
    seen_changes = set()
    unique_changes = []
    
    for change in changes:
        # Create a unique identifier for each change
        if change["type"] == "table_cell":
            change_id = f"table_{change.get('table_index')}_cell_{change.get('row')}_{change.get('col')}"
        elif change["type"] == "paragraph":
            change_id = f"para_{change.get('index')}"
        elif change["type"] == "image":
            change_id = f"image_{change.get('hash', '')}"
        elif change["type"] == "table_image":
            change_id = f"table_image_{change.get('table')}_{change.get('row')}_{change.get('col')}_{change.get('hash', '')}"
        else:
            change_id = f"{change['type']}_{hash(str(change))}"
        
        if change_id not in seen_changes:
            seen_changes.add(change_id)
            unique_changes.append(change)
    
    return unique_changes

def cluster_similar_changes(changes: List[Dict]) -> List[Dict]:
    """
    Cluster similar changes together for better organization.
    """
    if not changes:
        return changes
    
    # Extract features for clustering
    change_texts = []
    for change in changes:
        text = f"{change.get('type', '')} {change.get('status', '')} {change.get('change_detail', '')}"
        change_texts.append(text)
    
    # Use TF-IDF for text vectorization
    vectorizer = TfidfVectorizer(max_features=50, stop_words='english')
    try:
        X = vectorizer.fit_transform(change_texts)
        
        # Cluster using DBSCAN
        clustering = DBSCAN(eps=0.5, min_samples=2).fit(X.toarray())
        labels = clustering.labels_
        
        # Group changes by cluster
        clustered_changes = []
        for label in set(labels):
            if label == -1:  # Noise points, don't cluster
                continue
            cluster_indices = [i for i, l in enumerate(labels) if l == label]
            if len(cluster_indices) > 1:  # Only cluster if multiple similar changes
                cluster_changes = [changes[i] for i in cluster_indices]
                clustered_changes.append({
                    'type': 'change_cluster',
                    'cluster_id': f"cluster_{label}",
                    'changes': cluster_changes,
                    'summary': f"Group of {len(cluster_changes)} similar changes",
                    'representative_change': cluster_changes[0]
                })
        
        # Add non-clustered changes
        noise_indices = [i for i, l in enumerate(labels) if l == -1]
        for idx in noise_indices:
            clustered_changes.append(changes[idx])
        
        return clustered_changes
    except Exception as e:
        print(f"Clustering failed: {e}")
        return changes

# ================================
# EXCEL EXPORT FUNCTIONS
# ================================

def export_to_excel_enhanced(changes: List[Dict], stats: Dict[str, Any]) -> io.BytesIO:
    """
    Enhanced Excel export with interactive dashboard and advanced analytics.
    """
    return export_to_excel_enhanced_with_dashboard(changes, stats)

def export_to_excel_enhanced_with_dashboard(changes: List[Dict], stats: Dict[str, Any]) -> io.BytesIO:
    """
    Enhanced Excel export with interactive dashboard and advanced analytics.
    """
    wb = openpyxl.Workbook()
    
    # Remove default sheet if it exists
    if 'Sheet' in wb.sheetnames:
        del wb['Sheet']
    
    # Create Dashboard sheet
    dashboard_ws = wb.create_sheet("Dashboard")
    dashboard_ws.sheet_view.showGridLines = False
    
    # Title and summary
    dashboard_ws.merge_cells('A1:H1')
    title_cell = dashboard_ws['A1']
    title_cell.value = "ADVANCED DOCUMENT COMPARISON DASHBOARD"
    title_cell.font = Font(bold=True, size=18, color="366092")
    title_cell.alignment = Alignment(horizontal='center')
    
    # Key metrics
    metrics = [
        ["Document Comparison Analytics", ""],
        ["Total Changes Detected", stats["changes"]["total"]],
        ["Document Similarity Score", f"{stats.get('similarity_metrics', {}).get('overall_similarity', 0) * 100:.1f}%"],
        ["Change Density", f"{stats.get('change_density', 0) * 100:.1f}%"],
        ["Critical Changes", stats["changes"]["by_severity"].get(3, 0)],
        ["Major Changes", stats["changes"]["by_severity"].get(2, 0)],
        ["Processing Date", stats.get("processing_time", "N/A")],
    ]
    
    for i, (label, value) in enumerate(metrics, start=3):
        dashboard_ws[f'A{i}'] = label
        dashboard_ws[f'B{i}'] = value
        dashboard_ws[f'A{i}'].font = Font(bold=True)
    
    # Change distribution chart data
    dashboard_ws['A10'] = "Change Distribution by Type"
    type_data = []
    for change_type, count in stats["changes"]["by_type"].items():
        type_data.append((change_type.replace('_', ' ').title(), count))
    
    for i, (change_type, count) in enumerate(type_data, start=11):
        if i <= 20:  # Limit to 10 types for readability
            dashboard_ws[f'A{i}'] = change_type
            dashboard_ws[f'B{i}'] = count
    
    # Severity distribution
    dashboard_ws['D10'] = "Change Severity Distribution"
    severity_labels = {0: 'Info', 1: 'Minor', 2: 'Major', 3: 'Critical'}
    for i, (severity, count) in enumerate(stats["changes"]["by_severity"].items(), start=11):
        if i <= 15:  # Limit to 5 severity levels
            dashboard_ws[f'D{i}'] = severity_labels.get(severity, f"Level {severity}")
            dashboard_ws[f'E{i}'] = count
    
    # Set column widths
    dashboard_ws.column_dimensions['A'].width = 25
    dashboard_ws.column_dimensions['B'].width = 15
    dashboard_ws.column_dimensions['D'].width = 25
    dashboard_ws.column_dimensions['E'].width = 15
        
    # Create Summary sheet
    summary_ws = wb.create_sheet("Summary")
    summary_ws.sheet_view.showGridLines = False
    
    summary_ws.merge_cells('A1:D1')
    title_cell = summary_ws['A1']
    title_cell.value = "ENHANCED DOCUMENT COMPARISON REPORT"
    title_cell.font = Font(bold=True, size=16)
    title_cell.alignment = Alignment(horizontal='center')

    # Add summary information
    summary_data = [
        ["Document Statistics", ""],
        ["Total Paragraphs", stats["paragraphs"]],
        ["Total Tables Processed", stats["tables"]],
        ["Total Images", stats["images"]],
        ["Total Table Images", stats["table_images"]],
        ["Total Panels", stats["panels"]],
        ["", ""],
        ["Change Statistics", ""],
        ["Total Changes Detected", stats["changes"]["total"]],
        ["Change Density", f"{stats.get('change_density', 0) * 100:.2f}%"],
        ["Document Similarity", f"{stats.get('similarity_metrics', {}).get('overall_similarity', 0) * 100:.2f}%"],
    ]
    
    # Add change type breakdown
    summary_data.append(["", ""])
    summary_data.append(["Change Type Breakdown", ""])
    for change_type, count in stats["changes"]["by_type"].items():
        summary_data.append([f"  {change_type.replace('_', ' ').title()}", count])
    
    # Add severity breakdown
    summary_data.append(["", ""])
    summary_data.append(["Severity Breakdown", ""])
    for severity, count in stats["changes"]["by_severity"].items():
        severity_name = severity_labels.get(severity, f"Level {severity}")
        summary_data.append([f"  {severity_name}", count])
    
    for i, (label, value) in enumerate(summary_data, start=3):
        summary_ws[f'A{i}'] = label
        summary_ws[f'B{i}'] = value
        if not str(label).startswith('  '):  # Bold main categories
            summary_ws[f'A{i}'].font = Font(bold=True)
        
    # Set column widths for summary sheet
    summary_ws.column_dimensions['A'].width = 30
    summary_ws.column_dimensions['B'].width = 20
    
    
    # SEPARATE STRIKETHROUGH CHANGES FROM REGULAR CHANGES WITH DEDUPLICATION
    strikethrough_changes = []
    regular_changes = []
    seen_strikethrough_paragraphs = set()
    
    for change in changes:
        change_type = change.get('type', '')
        
        # Identify strikethrough changes
        if any(keyword in change_type.lower() for keyword in ['strikethrough', 'strike']):
            para_index = change.get('index')
            para_key = f"para_{para_index}" if para_index else None
            
            # Deduplicate by paragraph index
            if para_key and para_key not in seen_strikethrough_paragraphs:
                strikethrough_changes.append(change)
                seen_strikethrough_paragraphs.add(para_key)
            elif not para_key:  # For changes without paragraph index (table strikethrough)
                # Create unique key for table strikethrough changes
                table_key = f"table_{change.get('table_index')}_{change.get('row')}_{change.get('col')}"
                if table_key not in seen_strikethrough_paragraphs:
                    strikethrough_changes.append(change)
                    seen_strikethrough_paragraphs.add(table_key)
        else:
            regular_changes.append(change)
    
    print(f"DEBUG: After deduplication - Strikethrough: {len(strikethrough_changes)}, Regular: {len(regular_changes)}")
    
    # Create Strikethrough Changes sheet (ONLY for strikethrough changes)
    if strikethrough_changes:
        strike_ws = wb.create_sheet("Strikethrough Changes")
        strike_headers = [
            "Change #", "Content Type", "Location", "Status", "Severity",
            "Strikethrough Text", "Context", "Change Details"
        ]
        strike_ws.append(strike_headers)
        
        # Style strikethrough sheet headers
        for col in range(1, len(strike_headers) + 1):
            cell = strike_ws.cell(row=1, column=col)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
        # Process strikethrough changes
        strike_idx = 2
        seen_strike_details = set()
        
        for change in strikethrough_changes:
            # Create unique identifier for this change to avoid duplicates
            change_id = f"{change.get('type')}_{change.get('index')}_{change.get('table_index')}_{change.get('row')}_{change.get('col')}_{change.get('status')}"
            
            if change_id in seen_strike_details:
                continue
                
            seen_strike_details.add(change_id)
            
            # Set location based on change type
            if change.get("type") in ["paragraph_strikethrough", "paragraph_formatting_strikethrough"]:
                location = f"Paragraph {change.get('index', '')}"
                content_type = "Paragraph"
            elif change.get("type") in ["table_strikethrough", "table_cell_strikethrough"]:
                location = f"Table {change.get('table_index', '')}"
                if change.get("row"):
                    location += f", Cell ({change.get('row')},{change.get('col')})"
                content_type = "Table"
            else:
                location = "Unknown"
                content_type = "Unknown"
            
            # Extract strikethrough text
            strike_text = ""
            if change.get("strikethrough_text_original"):
                strike_text += f"Original: {change.get('strikethrough_text_original')}"
            if change.get("strikethrough_text_modified"):
                if strike_text:
                    strike_text += "\n"
                strike_text += f"Modified: {change.get('strikethrough_text_modified')}"
            if change.get("strikethrough_original") and not strike_text:
                strike_text = change.get("strikethrough_original", "")
            if change.get("strikethrough_modified") and not strike_text:
                strike_text = change.get("strikethrough_modified", "")
            
            # If no specific strikethrough text, use boolean indicators
            if not strike_text.strip():
                if change.get("strikethrough_original"):
                    strike_text = "Strikethrough in original document"
                elif change.get("strikethrough_modified"):
                    strike_text = "Strikethrough in modified document"
                    
            # Context information (truncated for readability)
            context = ""
            original_text = change.get("original", "")
            modified_text = change.get("modified", "")
            
            if original_text:
                context += f"Original: {original_text[:100]}{'...' if len(original_text) > 100 else ''}"
            if modified_text:
                if context:
                    context += "\n"
                context += f"Modified: {modified_text[:100]}{'...' if len(modified_text) > 100 else ''}"
            
            # Change details (clean up formatting changes)
            change_detail = change.get("change_detail", "")
            
            # Add severity information
            severity = change.get('severity', 0)
            severity_labels = {0: 'Info', 1: 'Minor', 2: 'Major', 3: 'Critical'}
            severity_text = severity_labels.get(severity, 'Unknown')
            
            strike_ws.append([
                strike_idx-1,
                content_type,
                location,
                change.get("status", "").replace("_", " ").title(),
                severity_text,
                strike_text,
                context,
                change_detail
            ])
            
            # Apply strikethrough formatting to the text cells
            text_cell = strike_ws.cell(row=strike_idx, column=6)  # Strikethrough Text column
            text_cell.font = Font(strike=True)
            
            # Color code based on severity
            severity = change.get('severity', 0)
            severity_colors = {
                0: "E6E6FA",  # Info - Lavender
                1: "90EE90",  # Minor - Light Green
                2: "FFFFE0",  # Major - Light Yellow
                3: "FFCCCB"   # Critical - Light Red
            }
            fill_color = severity_colors.get(severity, "FFFFFF")
            fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
            
            for col in range(1, len(strike_headers) + 1):
                strike_ws.cell(row=strike_idx, column=col).fill = fill
                strike_ws.cell(row=strike_idx, column=col).border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            
            strike_idx += 1
        
        # Set column widths for strikethrough sheet
        strike_widths = [8, 15, 25, 15, 10, 40, 50, 60]
        for i, width in enumerate(strike_widths, start=1):
            strike_ws.column_dimensions[get_column_letter(i)].width = width

    # Create Detailed Changes sheet for regular changes
    changes_ws = wb.create_sheet("Detailed Changes")
    headers = [
        "Change #", "Type", "Location", "Status", "Severity", "Category",
        "Original Content", "Modified Content", "Change Details"
    ]
    changes_ws.append(headers)
    
    # Style main changes sheet headers
    for col in range(1, len(headers) + 1):
        cell = changes_ws.cell(row=1, column=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    # Define color fills for different severity levels
    severity_colors = {
        0: "E6E6FA",  # Info - Lavender
        1: "90EE90",  # Minor - Light Green
        2: "FFFFE0",  # Major - Light Yellow
        3: "FFCCCB"   # Critical - Light Red
    }
    
    # Process regular changes
    change_idx = 2
    for change in regular_changes:
        # Set location based on change type
        location = "Unknown"
        if change["type"] == "paragraph":
            location = f"Paragraph {change.get('index', '')}"
        elif change["type"] == "table":
            location = f"Table {change.get('index', '')}"
        elif change["type"] == "table_cell":
            location = f"Table {change.get('table_index', '')}, Cell ({change.get('row','')},{change.get('col','')})"
        elif change["type"] == "table_row":
            location = f"Table {change.get('table_index', '')}, Row {change.get('row','')}"
        elif change["type"] == "image":
            location = "Document"
        elif change["type"] == "table_image":
            location = f"Table {change.get('table','')}, Cell ({change.get('row','')},{change.get('col','')})"
        elif change["type"] == "panel":
            location = f"Panel {change.get('index', '')}"
        
        # Handle images
        if change["type"] in ["image", "table_image"]:
            if change["status"] == "deleted" and change.get("old_img"):
                try:
                    # For images, we'll just note their presence in the text
                    original_text = f"[Image: {change.get('width','?')}x{change.get('height','?')}]"
                    modified_text = "[Image Deleted]"
                except Exception as e:
                    original_text = "[Image]"
                    modified_text = "[Image Deleted]"
            elif change["status"] == "added" and change.get("new_img"):
                try:
                    original_text = "[No Image]"
                    modified_text = f"[Image: {change.get('width','?')}x{change.get('height','?')}]"
                except Exception as e:
                    original_text = "[No Image]"
                    modified_text = "[Image Added]"
            else:
                original_text = change.get("original", "")
                modified_text = change.get("modified", "")
        else:
            original_text = change.get("original", "")
            modified_text = change.get("modified", "")
        
        # Build change details
        change_details = change.get("change_detail", "")
        
        # Add additional details based on change type
        if change["type"] == "paragraph":
            details_parts = []
            if change.get("double_spaces_original", 0) > 0 or change.get("double_spaces_modified", 0) > 0:
                details_parts.append(f"Double spaces: {change.get('double_spaces_original', 0)} → {change.get('double_spaces_modified', 0)}")
            if change.get("tbd_original") or change.get("tbd_modified"):
                details_parts.append(f"TBD content: {change.get('tbd_original')} → {change.get('tbd_modified')}")
            if change.get("code_names_original") or change.get("code_names_modified"):
                details_parts.append(f"Code names: {change.get('code_names_original')} → {change.get('code_names_modified')}")
            if change.get("urls_original") or change.get("urls_modified"):
                details_parts.append(f"URLs: {change.get('urls_original')} → {change.get('urls_modified')}")
            
            if details_parts:
                change_details += "\n" + "\n".join(details_parts) if change_details else "\n".join(details_parts)
                
        elif change["type"] == "paragraph_formatting":
            formatting_details = []
            for fmt_change in change.get("formatting_changes", []):
                formatting_details.append(f"'{fmt_change['text']}': {', '.join(fmt_change['changes'])}")
            if formatting_details:
                change_details += "\nFormatting changes:\n" + "\n".join(formatting_details)
        
        # Get severity and category
        severity = change.get('severity', 0)
        severity_labels = {0: 'Info', 1: 'Minor', 2: 'Major', 3: 'Critical'}
        severity_text = severity_labels.get(severity, 'Unknown')
        category = change.get('category', 'Unknown').title()
        
        # Write row to changes sheet
        changes_ws.cell(row=change_idx, column=1, value=change_idx-1)  # Change #
        changes_ws.cell(row=change_idx, column=2, value=change["type"].replace("_", " ").title())
        changes_ws.cell(row=change_idx, column=3, value=location)
        changes_ws.cell(row=change_idx, column=4, value=change["status"].replace("_", " ").title())
        changes_ws.cell(row=change_idx, column=5, value=severity_text)
        changes_ws.cell(row=change_idx, column=6, value=category)
        changes_ws.cell(row=change_idx, column=7, value=str(original_text)[:500])  # Limit length
        changes_ws.cell(row=change_idx, column=8, value=str(modified_text)[:500])  # Limit length
        changes_ws.cell(row=change_idx, column=9, value=str(change_details)[:1000])  # Limit length
            
        # Apply fill based on severity
        fill_color = severity_colors.get(severity, "FFFFFF")
        fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        
        # Apply styling to all cells in the row
        for col in range(1, len(headers) + 1):
            cell = changes_ws.cell(row=change_idx, column=col)
            cell.fill = fill
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            # Enable text wrapping for content cells
            if col >= 7:  # Content columns
                cell.alignment = Alignment(wrap_text=True, vertical='top')
        
        change_idx += 1

    # Set column widths for changes sheet
    column_widths = [8, 15, 25, 12, 10, 12, 40, 40, 60]
    for i, width in enumerate(column_widths, start=1):
        changes_ws.column_dimensions[get_column_letter(i)].width = width

    # Create Advanced Analytics sheet
    analytics_ws = wb.create_sheet("Advanced Analytics")
    analytics_headers = ["Metric", "Value", "Description"]
    analytics_ws.append(analytics_headers)
    
    similarity_metrics = stats.get('similarity_metrics', {})
    analytics_data = [
        ["Document Similarity", f"{similarity_metrics.get('overall_similarity', 0) * 100:.2f}%", "Overall similarity score"],
        ["Text Similarity", f"{similarity_metrics.get('text_similarity', 0) * 100:.2f}%", "Text content similarity"],
        ["Structural Similarity", f"{similarity_metrics.get('structural_similarity', 0) * 100:.2f}%", "Document structure similarity"],
        ["Content Preservation", f"{similarity_metrics.get('content_preservation', 0) * 100:.2f}%", "Percentage of preserved content"],
        ["Change Density", f"{stats.get('change_density', 0) * 100:.2f}%", "Changes per document element"],
        ["Total Elements", stats.get('total_elements', 0), "Total paragraphs, tables, and panels"],
    ]
    
    for row_data in analytics_data:
        analytics_ws.append(row_data)
    
    # Style analytics sheet
    for col in range(1, len(analytics_headers) + 1):
        cell = analytics_ws.cell(row=1, column=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    # Set column widths for analytics sheet
    analytics_ws.column_dimensions['A'].width = 25
    analytics_ws.column_dimensions['B'].width = 20
    analytics_ws.column_dimensions['C'].width = 50
    
    # Set active sheet to Dashboard
    wb.active = dashboard_ws
        
    # Save to bytes buffer
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    print(f"DEBUG: Excel export completed - {len(regular_changes)} regular changes, {len(strikethrough_changes)} strikethrough changes")
    return output

# ================================
# NEW: Batch Processing Functions
# ================================

def process_batch_comparison(file_pairs: List[Tuple], batch_id: str) -> Dict:
    """
    Process multiple document comparisons in batch mode.
    """
    results = {
        'batch_id': batch_id,
        'start_time': datetime.now().isoformat(),
        'total_pairs': len(file_pairs),
        'completed': 0,
        'results': [],
        'errors': []
    }
    
    def process_single_pair(pair_idx, file1, file2, pass1, pass2):
        try:
            comparison_id = f"{batch_id}_{pair_idx}"
            update_progress(comparison_id, "Starting", 0)
            
            doc1 = load_docx(file1, pass1)
            doc2 = load_docx(file2, pass2)
            
            changes, stats = compare_docx_enhanced(doc1, doc2, comparison_id)
            
            return {
                'pair_index': pair_idx,
                'file1_name': getattr(file1, 'filename', 'Unknown'),
                'file2_name': getattr(file2, 'filename', 'Unknown'),
                'changes_count': len(changes),
                'similarity': stats.get('similarity_metrics', {}).get('overall_similarity', 0),
                'success': True
            }
        except Exception as e:
            return {
                'pair_index': pair_idx,
                'error': str(e),
                'success': False
            }
    
    # Process pairs with threading
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = []
        for idx, (file1, file2, pass1, pass2) in enumerate(file_pairs):
            future = executor.submit(process_single_pair, idx, file1, file2, pass1, pass2)
            futures.append(future)
        
        for future in as_completed(futures):
            result = future.result()
            if result['success']:
                results['results'].append(result)
            else:
                results['errors'].append(result)
            results['completed'] += 1
    
    results['end_time'] = datetime.now().isoformat()
    results['success_rate'] = len(results['results']) / len(file_pairs) if file_pairs else 0
    
    return results

# ================================
# NEW: Advanced Filtering System
# ================================

class ChangeFilter:
    """Advanced change filtering system"""
    
    def __init__(self):
        self.filters = {}
    
    def apply_filters(self, changes: List[Dict], filters: Dict) -> List[Dict]:
        """Apply multiple filters to changes"""
        filtered_changes = changes
        
        if filters.get('severity_min') is not None:
            filtered_changes = [c for c in filtered_changes if c.get('severity', 0) >= filters['severity_min']]
        
        if filters.get('severity_max') is not None:
            filtered_changes = [c for c in filtered_changes if c.get('severity', 0) <= filters['severity_max']]
        
        if filters.get('types'):
            filtered_changes = [c for c in filtered_changes if c.get('type') in filters['types']]
        
        if filters.get('categories'):
            filtered_changes = [c for c in filtered_changes if c.get('category') in filters['categories']]
        
        if filters.get('search_text'):
            search_text = filters['search_text'].lower()
            filtered_changes = [
                c for c in filtered_changes 
                if any(search_text in str(val).lower() for val in c.values() if isinstance(val, str))
            ]
        
        return filtered_changes
    
    def get_filter_summary(self, original_changes: List[Dict], filtered_changes: List[Dict]) -> Dict:
        """Get summary of filtering results"""
        return {
            'original_count': len(original_changes),
            'filtered_count': len(filtered_changes),
            'removed_count': len(original_changes) - len(filtered_changes),
            'filter_efficiency': round((len(original_changes) - len(filtered_changes)) / len(original_changes) * 100, 1)
        }

# ================================
# Enhanced Flask Routes
# ================================

@app.route("/progress/<comparison_id>")
def get_progress(comparison_id: str):
    """Get progress updates for long-running comparisons"""
    progress = progress_data.get(comparison_id, {'stage': 'Unknown', 'progress': 0, 'total': 100})
    return jsonify(progress)

@app.route("/batch_compare", methods=["POST"])
def batch_compare():
    """Handle batch document comparisons"""
    try:
        files = request.files.getlist("documents[]")
        passwords = request.form.getlist("passwords[]")
        
        if len(files) % 2 != 0:
            return jsonify({"error": "Please upload pairs of documents"}), 400
        
        # Group files into pairs
        file_pairs = []
        for i in range(0, len(files), 2):
            file1 = files[i]
            file2 = files[i + 1]
            pass1 = passwords[i] if i < len(passwords) else None
            pass2 = passwords[i + 1] if i + 1 < len(passwords) else None
            
            file_pairs.append((file1, file2, pass1, pass2))
        
        batch_id = f"batch_{int(time.time())}"
        
        # Start batch processing in background thread
        def process_batch():
            results = process_batch_comparison(file_pairs, batch_id)
            # Save results to cache
            with open(f"{COMPARISON_CACHE}/{batch_id}.json", 'w') as f:
                json.dump(results, f)
        
        thread = threading.Thread(target=process_batch)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            "batch_id": batch_id,
            "message": f"Batch processing started for {len(file_pairs)} pairs",
            "status_url": f"/batch_status/{batch_id}"
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/batch_status/<batch_id>")
def batch_status(batch_id: str):
    """Get status of batch processing"""
    try:
        results_path = f"{COMPARISON_CACHE}/{batch_id}.json"
        if os.path.exists(results_path):
            with open(results_path, 'r') as f:
                results = json.load(f)
            return jsonify(results)
        else:
            return jsonify({"status": "processing", "message": "Batch processing in progress"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/filter_changes", methods=["POST"])
def filter_changes():
    """Apply filters to changes"""
    try:
        changes = request.json.get('changes', [])
        filters = request.json.get('filters', {})
        
        change_filter = ChangeFilter()
        filtered_changes = change_filter.apply_filters(changes, filters)
        summary = change_filter.get_filter_summary(changes, filtered_changes)
        
        return jsonify({
            "filtered_changes": filtered_changes,
            "summary": summary
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/visualize_changes", methods=["POST"])
def visualize_changes():
    """Generate visualization for changes"""
    try:
        changes = request.json.get('changes', [])
        stats = request.json.get('stats', {})
        
        # Create severity distribution plot
        severity_counts = Counter([change.get('severity', 0) for change in changes])
        
        plt.figure(figsize=(10, 6))
        plt.bar(severity_counts.keys(), severity_counts.values())
        plt.xlabel('Severity Level')
        plt.ylabel('Number of Changes')
        plt.title('Change Severity Distribution')
        
        # Save plot to bytes
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        img_data = base64.b64encode(img_buffer.getvalue()).decode()
        
        return jsonify({
            "visualization": f"data:image/png;base64,{img_data}",
            "summary": {
                "total_changes": len(changes),
                "severity_distribution": dict(severity_counts)
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ================================
# Enhanced Main Route
# ================================

@app.route("/", methods=["GET", "POST"])
def index():
    """
    Main route for document comparison interface.
    Handles both GET requests (page display) and POST requests (document processing).
    """
    if request.method == "POST":
        is_download_request = request.form.get('download') == 'true'
        
        file1 = request.files.get("doc1")
        file2 = request.files.get("doc2")
        pass1 = request.form.get("password1", "").strip() or None
        pass2 = request.form.get("password2", "").strip() or None

        if not file1 or not file2:
            error_msg = "Please upload both documents."
            if is_download_request or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({"error": error_msg}), 400
            flash(f"⚠️ {error_msg}")
            return render_template("index.html")
            
        try:
            # Reset stream positions
            file1.stream.seek(0)
            file2.stream.seek(0)
            
            # Check if documents need passwords
            needs_pass1 = is_encrypted(file1.stream)
            needs_pass2 = is_encrypted(file2.stream)
            
            # Reset streams again after encryption check
            file1.stream.seek(0)
            file2.stream.seek(0)
            
            if needs_pass1 and not pass1:
                error_msg = "Document 1 is encrypted but no password was provided."
                if is_download_request or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({"error": error_msg}), 400
                flash(f"⚠️ {error_msg}")
                return render_template("index.html")
                
            if needs_pass2 and not pass2:
                error_msg = "Document 2 is encrypted but no password was provided."
                if is_download_request or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({"error": error_msg}), 400
                flash(f"⚠️ {error_msg}")
                return render_template("index.html")

            # Load documents
            doc1 = load_docx(file1.stream, pass1)
            doc2 = load_docx(file2.stream, pass2)
            
            comparison_id = f"comp_{int(time.time())}"

            # Compare documents with enhanced comparison
            changes, stats = compare_docx_enhanced(doc1, doc2, comparison_id)
            
            # Add clustering if requested
            if request.form.get('enable_clustering'):
                changes = cluster_similar_changes(changes)

            # For download requests, return the Excel file
            if is_download_request:
                excel_file = export_to_excel_enhanced(changes, stats)
                
                # Generate filename with timestamp
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"document_comparison_{timestamp}.xlsx"
                return send_file(
                    excel_file,
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            # For AJAX requests, return JSON
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                # Convert changes to serializable format
                serializable_changes = []
                for change in changes:
                    serializable_change = change.copy()
                    # Remove image bytes as they can't be serialized to JSON
                    serializable_change.pop('old_img', None)
                    serializable_change.pop('new_img', None)
                    serializable_changes.append(serializable_change)
                
                return jsonify({
                    "changes": serializable_changes,
                    "stats": stats
                })
            else:
                # For regular form submission, show results on page
                return render_template("index.html", changes=changes, stats=stats)

        except Exception as e:
            error_msg = f"Error processing documents: {str(e)}"
            if is_download_request or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({"error": error_msg}), 400
            else:
                flash(f"❌ {error_msg}")
                return render_template("index.html")

    return render_template("index.html")

# ================================
# NEW: API Routes for Integration
# ================================

@app.route("/api/compare", methods=["POST"])
def api_compare():
    """REST API endpoint for document comparison"""
    try:
        file1 = request.files.get('document1')
        file2 = request.files.get('document2')
        pass1 = request.form.get('password1')
        pass2 = request.form.get('password2')
        
        if not file1 or not file2:
            return jsonify({"error": "Both documents are required"}), 400
        
        doc1 = load_docx(file1.stream, pass1)
        doc2 = load_docx(file2.stream, pass2)
        
        changes, stats = compare_docx_enhanced(doc1, doc2)
        
        return jsonify({
            "success": True,
            "changes_count": len(changes),
            "similarity_score": stats.get('similarity_metrics', {}).get('overall_similarity', 0),
            "changes_preview": changes[:10],
            "stats": stats
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/health")
def api_health():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.0.0"
    })

# ================================
# Enhanced Template Context
# ================================

@app.context_processor
def utility_processor():
    """Add utility functions to template context"""
    def format_severity(severity):
        severity_labels = {0: 'Info', 1: 'Minor', 2: 'Major', 3: 'Critical'}
        return severity_labels.get(severity, 'Unknown')
    
    def get_severity_color(severity):
        severity_colors = {0: 'info', 1: 'success', 2: 'warning', 3: 'danger'}
        return severity_colors.get(severity, 'secondary')
    
    return {
        'format_severity': format_severity,
        'get_severity_color': get_severity_color,
        'current_year': datetime.now().year
    }

# ================================
# Error Handlers
# ================================

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large. Maximum size is 100MB."}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({"error": "Internal server error. Please try again."}), 500

# ================================
# APPLICATION ENTRY POINT
# ================================

if __name__ == "__main__":
    """
    Enhanced main entry point with additional configuration.
    """
    # Clean up old cache files
    for file in os.listdir(COMPARISON_CACHE):
        if file.endswith('.json'):
            file_path = os.path.join(COMPARISON_CACHE, file)
            # Delete files older than 24 hours
            if os.path.getctime(file_path) < time.time() - 24 * 3600:
                os.remove(file_path)
    
    app.run(debug=True, threaded=True)